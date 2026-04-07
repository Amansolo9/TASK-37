"""File upload, storage, validation, watermark, and download service."""

import hashlib
import hmac
import os
import time
import uuid
from datetime import datetime, timedelta
from pathlib import Path

from flask import current_app
from PIL import Image, ImageDraw, ImageFont
from app.extensions import db
from app.models.files import Attachment, FileDownloadAudit
from app.services.audit_service import log_action

ALLOWED_EXTENSIONS = {"pdf", "jpg", "jpeg", "png", "docx"}
BLOCKED_EXTENSIONS = {
    "exe", "bat", "cmd", "com", "msi", "scr", "pif", "vbs", "js", "wsf",
    "ps1", "sh", "cgi", "jar", "py", "rb", "pl",
}
# Strict mapping: extension -> set of exact allowed MIME types
MIME_MAP = {
    "pdf": {"application/pdf"},
    "jpg": {"image/jpeg"},
    "jpeg": {"image/jpeg"},
    "png": {"image/png"},
    "docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
}


# Magic byte signatures for content sniffing
_MAGIC_SIGNATURES = {
    b"%PDF": "application/pdf",
    b"\xff\xd8\xff": "image/jpeg",
    b"\x89PNG": "image/png",
    b"PK\x03\x04": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _sniff_mime(header_bytes, ext):
    """Sniff MIME type from file header bytes. Returns detected type or None."""
    if not header_bytes:
        return None
    for sig, mime in _MAGIC_SIGNATURES.items():
        if header_bytes[:len(sig)] == sig:
            return mime
    return None


def validate_upload(file_storage):
    """Validate file extension, MIME, and size. Returns (ok, error_message)."""
    filename = file_storage.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext in BLOCKED_EXTENSIONS:
        return False, f"File type .{ext} is blocked."
    if ext not in ALLOWED_EXTENSIONS:
        return False, f"Only {', '.join(ALLOWED_EXTENSIONS)} files are allowed."

    content_type = file_storage.content_type or ""
    allowed_mimes = MIME_MAP.get(ext, set())
    if allowed_mimes and content_type:
        if content_type not in allowed_mimes:
            return False, f"MIME type '{content_type}' is not allowed for .{ext} files. Expected: {', '.join(allowed_mimes)}."
    elif not content_type:
        return False, "MIME type is required for upload validation."

    # Server-side MIME sniffing: read first bytes to detect real type
    file_storage.seek(0)
    header = file_storage.read(16)
    file_storage.seek(0)
    sniffed_type = _sniff_mime(header, ext)
    if sniffed_type and sniffed_type not in allowed_mimes:
        return False, f"File content detected as '{sniffed_type}' which does not match .{ext}."

    file_storage.seek(0, 2)
    size = file_storage.tell()
    file_storage.seek(0)
    max_bytes = current_app.config["MAX_UPLOAD_MB"] * 1024 * 1024
    if size > max_bytes:
        return False, f"File exceeds {current_app.config['MAX_UPLOAD_MB']} MB limit."

    return True, None


def compute_sha256(file_storage):
    sha = hashlib.sha256()
    while True:
        chunk = file_storage.read(8192)
        if not chunk:
            break
        sha.update(chunk)
    file_storage.seek(0)
    return sha.hexdigest()


def save_upload(file_storage, user_id, owner_type=None, owner_id=None):
    ok, error = validate_upload(file_storage)
    if not ok:
        raise ValueError(error)

    sha = compute_sha256(file_storage)
    filename = file_storage.filename or "unnamed"
    ext = filename.rsplit(".", 1)[-1].lower()

    duplicate = Attachment.query.filter_by(sha256=sha, deleted_at=None).first()

    stored_name = f"{uuid.uuid4().hex}.{ext}"
    upload_dir = Path(current_app.config["STORAGE_ROOT"]) / "uploads"
    upload_dir.mkdir(parents=True, exist_ok=True)
    storage_path = upload_dir / stored_name
    file_storage.save(str(storage_path))

    file_storage.seek(0, 2)
    size = file_storage.tell()

    att = Attachment(
        owner_type=owner_type, owner_id=owner_id,
        original_filename=filename, stored_filename=stored_name,
        storage_path=str(storage_path), file_ext=ext,
        mime_type=next(iter(MIME_MAP.get(ext, set())), "application/octet-stream"),
        size_bytes=size, sha256=sha,
        duplicate_of_id=duplicate.id if duplicate else None,
        watermark_on_download=current_app.config.get("WATERMARK_DEFAULT_ENABLED", False),
        uploaded_by=user_id,
    )
    db.session.add(att)
    db.session.commit()
    log_action(user_id, "file_uploaded", "attachment", att.id,
              {"filename": filename, "size": size, "duplicate": duplicate is not None})
    return att


def generate_signed_url(attachment_id, user_id, api=False):
    ttl = current_app.config["DOWNLOAD_URL_TTL_SECONDS"]
    expires = int(time.time()) + ttl
    payload = f"{attachment_id}:{user_id}:{expires}"
    sig = hmac.new(
        current_app.config["SECRET_KEY"].encode(),
        payload.encode(), hashlib.sha256
    ).hexdigest()
    if api:
        return f"/api/v1/files/{attachment_id}/download?sig={sig}&expires={expires}&uid={user_id}"
    return f"/files/{attachment_id}/download?sig={sig}&expires={expires}&uid={user_id}"


def verify_signed_url(attachment_id, sig, expires, user_id):
    """Verify a signed download URL. Fails closed on any malformed input."""
    try:
        expires_int = int(expires)
    except (ValueError, TypeError):
        return False
    if not sig or not user_id:
        return False
    if time.time() > expires_int:
        return False
    try:
        payload = f"{attachment_id}:{user_id}:{expires_int}"
        expected = hmac.new(
            current_app.config["SECRET_KEY"].encode(),
            payload.encode(), hashlib.sha256
        ).hexdigest()
        return hmac.compare_digest(str(sig), expected)
    except Exception:
        return False


def get_download_path(attachment_id, user_id, apply_watermark=False):
    att = db.session.get(Attachment, attachment_id)
    if not att or att.deleted_at:
        return None, None

    audit = FileDownloadAudit(
        attachment_id=att.id, user_id=user_id,
        watermark_applied=apply_watermark,
    )
    db.session.add(audit)
    db.session.commit()

    if apply_watermark and att.watermark_on_download:
        try:
            return _apply_watermark(att, user_id), att.original_filename
        except Exception as e:
            import logging
            logging.getLogger(__name__).warning(
                "Watermark failed for attachment %s: %s", att.id, e
            )

    return att.storage_path, att.original_filename


def _apply_watermark(att, user_id):
    """Generate a temporary watermarked copy for download.

    Supported media types:
    - JPG/JPEG/PNG: visible text overlay via Pillow
    - PDF: text stamp on each page via pypdf + reportlab (if available)
    - DOCX: not supported (returns original file)
    """
    if att.file_ext in ("jpg", "jpeg", "png"):
        return _watermark_image(att, user_id)
    if att.file_ext == "pdf":
        return _watermark_pdf(att, user_id)
    if att.file_ext == "docx":
        return _watermark_docx(att, user_id)
    import logging
    logging.getLogger(__name__).info(
        "Watermark not supported for .%s files (attachment %s)", att.file_ext, att.id
    )
    return att.storage_path


def _watermark_pdf(att, user_id):
    """Add a visible text watermark to each page of a PDF.

    Creates a transparent overlay page with diagonal watermark text using
    the PDF content stream directly (no reportlab dependency), then merges
    it onto every page of the source document.
    """
    try:
        from pypdf import PdfReader, PdfWriter, PageObject, Transformation
        from pypdf.generic import (
            ArrayObject, DictionaryObject, NameObject,
            NumberObject, TextStringObject, ContentStream,
        )
        import io

        watermark_text = f"GreenCycle - User {user_id} - {datetime.utcnow().strftime('%Y-%m-%d')}"

        # Build a single-page PDF with the watermark text as a content stream
        wm_stream = (
            b"q\n"
            b"0.85 0.85 0.85 rg\n"            # light-gray fill
            b"BT\n"
            b"/F1 36 Tf\n"                     # 36pt font
            b"1 0 0.4 1 100 300 Tm\n"          # rotated ~22 deg, positioned mid-page
            + watermark_text.encode("latin-1", errors="replace")
            + b" Tj\n"
            b"ET\n"
            b"Q\n"
        )

        reader = PdfReader(att.storage_path)
        writer = PdfWriter()

        for page in reader.pages:
            # Ensure the page has a Helvetica font resource for /F1
            resources = page.get("/Resources")
            if resources is None:
                page[NameObject("/Resources")] = DictionaryObject()
                resources = page["/Resources"]
            fonts = resources.get("/Font")
            if fonts is None:
                resources[NameObject("/Font")] = DictionaryObject()
                fonts = resources["/Font"]
            if "/F1" not in fonts:
                fonts[NameObject("/F1")] = DictionaryObject({
                    NameObject("/Type"): NameObject("/Font"),
                    NameObject("/Subtype"): NameObject("/Type1"),
                    NameObject("/BaseFont"): NameObject("/Helvetica"),
                })

            # Append the watermark content stream to the page
            existing = page.get("/Contents")
            if existing is not None:
                if not isinstance(existing, ArrayObject):
                    existing = ArrayObject([existing])
            else:
                existing = ArrayObject()
            wm_ref = writer._add_object(
                ContentStream(io.BytesIO(wm_stream), None)
            )
            existing.append(wm_ref)
            page[NameObject("/Contents")] = existing
            writer.add_page(page)

        tmp_dir = Path(current_app.config["STORAGE_ROOT"]) / "tmp"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        tmp_path = tmp_dir / f"wm_{uuid.uuid4().hex}.pdf"
        with open(tmp_path, "wb") as f:
            writer.write(f)
        return str(tmp_path)
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning("PDF watermark failed for attachment %s: %s", att.id, e)
        return att.storage_path


def _watermark_docx(att, user_id):
    """Add a visible watermark to a DOCX file via header text."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor

        doc = Document(att.storage_path)
        watermark_text = f"GreenCycle - User {user_id} - {datetime.utcnow().strftime('%Y-%m-%d')}"

        # Add watermark to every section's header
        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            para = header.add_paragraph()
            run = para.add_run(watermark_text)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(180, 180, 180)
            run.font.italic = True

        tmp_dir = Path(current_app.config["STORAGE_ROOT"]) / "tmp"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        tmp_path = tmp_dir / f"wm_{uuid.uuid4().hex}.docx"
        doc.save(str(tmp_path))
        return str(tmp_path)
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning("DOCX watermark failed for attachment %s: %s", att.id, e)
        return att.storage_path


def _watermark_image(att, user_id):
    img = Image.open(att.storage_path)
    draw = ImageDraw.Draw(img)
    text = f"GreenCycle - User {user_id} - {datetime.utcnow().strftime('%Y-%m-%d')}"
    try:
        font = ImageFont.truetype("arial.ttf", 20)
    except Exception:
        font = ImageFont.load_default()
    draw.text((10, 10), text, fill=(128, 128, 128, 128), font=font)
    tmp_dir = Path(current_app.config["STORAGE_ROOT"]) / "tmp"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    tmp_path = tmp_dir / f"wm_{uuid.uuid4().hex}.{att.file_ext}"
    img.save(str(tmp_path))
    return str(tmp_path)


def archive_old_attachments():
    cutoff = datetime.utcnow() - timedelta(days=current_app.config["ATTACHMENT_ARCHIVE_DAYS"])
    to_archive = Attachment.query.filter(
        Attachment.uploaded_at < cutoff,
        Attachment.archived_at.is_(None),
        Attachment.deleted_at.is_(None),
    ).all()
    archive_dir = Path(current_app.config["STORAGE_ROOT"]) / "archive"
    archive_dir.mkdir(parents=True, exist_ok=True)
    count = 0
    for att in to_archive:
        src = Path(att.storage_path)
        if src.exists():
            dest = archive_dir / att.stored_filename
            src.rename(dest)
            att.storage_path = str(dest)
        att.archived_at = datetime.utcnow()
        count += 1
    if count:
        db.session.commit()
    return count


def purge_expired_attachments():
    if not current_app.config.get("ENABLE_FILE_PURGE"):
        return 0
    cutoff = datetime.utcnow() - timedelta(days=current_app.config["ATTACHMENT_PURGE_YEARS"] * 365)
    to_purge = Attachment.query.filter(
        Attachment.uploaded_at < cutoff,
        Attachment.deleted_at.is_(None),
    ).all()
    count = 0
    for att in to_purge:
        path = Path(att.storage_path)
        if path.exists():
            path.unlink()
        att.deleted_at = datetime.utcnow()
        count += 1
    if count:
        db.session.commit()
    return count
